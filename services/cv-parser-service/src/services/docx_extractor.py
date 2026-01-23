import docx
import zipfile
import xml.etree.ElementTree as ET
import logging
from typing import Optional
import io

logger = logging.getLogger(__name__)

class DOCXExtractor:
    """Extracteur de texte depuis les fichiers DOCX"""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Optional[str]:
        """
        Extraire le texte d'un fichier DOCX - ancienne méthode pour compatibilité
        """
        try:
            with open(file_path, 'rb') as f:
                docx_bytes = f.read()
            return DOCXExtractor.extract_text_from_docx_bytes(docx_bytes)
                
        except Exception as e:
            logger.error(f"❌ Erreur lors de l'extraction DOCX: {e}")
            raise ValueError(f"Impossible d'extraire le texte du DOCX: {e}")
    
    @staticmethod
    def extract_text_from_docx_bytes(docx_bytes: bytes) -> Optional[str]:
        """
        Extraire le texte depuis des bytes DOCX (100% en mémoire)
        """
        try:
            # Utiliser python-docx avec des bytes
            doc = docx.Document(io.BytesIO(docx_bytes))
            text_parts = []
            
            # Extraire les paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n\n".join(text_parts)
            
            if text.strip():
                logger.info(f"✅ DOCX extrait depuis bytes avec python-docx ({len(text)} caractères)")
                return text.strip()
            else:
                # Essayer la méthode XML si python-docx ne retourne rien
                return DOCXExtractor._extract_from_xml_bytes(docx_bytes)
            
        except Exception as e:
            logger.warning(f"python-docx échoué avec bytes: {e}, essai méthode XML")
            try:
                return DOCXExtractor._extract_from_xml_bytes(docx_bytes)
            except Exception as e2:
                logger.error(f"Erreur extraction DOCX bytes: {e2}")
                raise ValueError(f"Impossible d'extraire le texte du DOCX: {e2}")
    
    @staticmethod
    def _extract_from_xml(file_path: str) -> str:
        """Extraire le texte directement depuis le XML du DOCX"""
        try:
            with open(file_path, 'rb') as f:
                docx_bytes = f.read()
            return DOCXExtractor._extract_from_xml_bytes(docx_bytes)
                
        except Exception as e:
            logger.error(f"Erreur extraction XML DOCX: {e}")
            raise
    
    @staticmethod
    def _extract_from_xml_bytes(docx_bytes: bytes) -> str:
        """Extraire le texte depuis des bytes DOCX via XML (100% en mémoire)"""
        try:
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
                document_xml = docx_zip.read('word/document.xml')
                root = ET.fromstring(document_xml)
                
                # Namespace
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Extraire tout le texte
                text_parts = []
                
                # Chercher tous les éléments de texte
                for t in root.iter():
                    if t.text and t.text.strip():
                        text_parts.append(t.text.strip())
                
                # Chercher spécifiquement dans les paragraphes
                for paragraph in root.findall('.//w:p', ns):
                    paragraph_text = []
                    for elem in paragraph.iter():
                        if elem.tag.endswith('}t'):  # Élément texte
                            if elem.text and elem.text.strip():
                                paragraph_text.append(elem.text.strip())
                    if paragraph_text:
                        text_parts.append(' '.join(paragraph_text))
                
                text = '\n\n'.join(text_parts)
                logger.info(f"✅ DOCX extrait depuis bytes via XML ({len(text)} caractères)")
                return text.strip()
                
        except Exception as e:
            logger.error(f"Erreur extraction XML bytes DOCX: {e}")
            raise